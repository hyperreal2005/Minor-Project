"""Formatting-preserving edit helpers for the ForgetCheck master reference."""
import copy
import docx
from docx.table import Table
from docx.text.paragraph import Paragraph


def blocks(doc):
    """Body children as (kind, obj) in document order — 'P' or 'T'."""
    out = []
    for child in doc.element.body.iterchildren():
        if child.tag.endswith('}p'):
            out.append(('P', Paragraph(child, doc)))
        elif child.tag.endswith('}tbl'):
            out.append(('T', Table(child, doc)))
    return out


def set_text(para, text):
    """Replace a paragraph's text, keeping the first run's character formatting."""
    runs = para.runs
    if not runs:
        para.add_run(text)
        return para
    runs[0].text = text
    for r in runs[1:]:
        r._element.getparent().remove(r._element)
    return para


def cell_text(cell, text):
    """Replace a table cell's text, keeping formatting and dropping extra paragraphs."""
    ps = cell.paragraphs
    set_text(ps[0], text)
    for p in ps[1:]:
        p._element.getparent().remove(p._element)
    return cell


def row_values(table, r, values):
    """Set a whole row from a list, skipping None entries."""
    for c, v in enumerate(values):
        if v is not None:
            cell_text(table.rows[r].cells[c], v)


def append_row(table, values, template_row=-1):
    """Append a row cloned from `template_row` and fill it."""
    src = table.rows[template_row]._tr
    new = copy.deepcopy(src)
    table.rows[-1]._tr.addnext(new)
    idx = len(table.rows) - 1
    row_values(table, idx, values)
    return idx


def insert_para_after(para, text, style=None):
    """Insert a new paragraph after `para`, cloning its formatting."""
    new_p = copy.deepcopy(para._element)
    para._element.addnext(new_p)
    np = Paragraph(new_p, para._parent)
    set_text(np, text)
    if style is not None:
        np.style = style
    return np


def insert_paras_after(para, items):
    """Insert several (text, style) pairs after `para`, in order. Returns the last."""
    cur = para
    for text, style in items:
        cur = insert_para_after(cur, text, style)
    return cur


def find_para(doc, needle, start=0):
    """Index of the first block paragraph containing `needle`."""
    bl = blocks(doc)
    for i in range(start, len(bl)):
        k, o = bl[i]
        if k == 'P' and needle in o.text:
            return i
    raise KeyError(f"paragraph not found: {needle!r}")


def find_table_by_header(doc, first_cell_text):
    """Return the first table whose top-left cell matches."""
    for k, o in blocks(doc):
        if k == 'T' and o.rows[0].cells[0].text.strip() == first_cell_text:
            return o
    raise KeyError(f"table not found with header cell: {first_cell_text!r}")


def find_callout(doc, label):
    """Return the 1x1 callout table whose first paragraph is `label`."""
    for k, o in blocks(doc):
        if k == 'T' and len(o.rows) == 1 and len(o.columns) == 1:
            ps = o.rows[0].cells[0].paragraphs
            if ps and ps[0].text.strip() == label:
                return o
    raise KeyError(f"callout not found: {label!r}")


def set_callout(doc, label, body, new_label=None):
    """Replace a callout's body paragraph, optionally relabelling it."""
    tbl = find_callout(doc, label)
    ps = tbl.rows[0].cells[0].paragraphs
    if new_label:
        set_text(ps[0], new_label)
    set_text(ps[1], body)
    for p in ps[2:]:
        p._element.getparent().remove(p._element)
    return tbl


def strip_hyperlinks(cell):
    """Remove w:hyperlink elements from a cell (they survive run-level edits)."""
    ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
    for p in cell.paragraphs:
        for hl in p._element.findall(ns + 'hyperlink'):
            p._element.remove(hl)
    return cell


def cell_text_clean(cell, text):
    """cell_text, but also drops hyperlinks so stale links don't linger."""
    strip_hyperlinks(cell)
    ps = cell.paragraphs
    if not ps[0].runs:
        ps[0].add_run('')
    set_text(ps[0], text)
    for p in ps[1:]:
        p._element.getparent().remove(p._element)
    return cell
